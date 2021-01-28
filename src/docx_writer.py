import docx
import docx.opc.constants as constants
import win32com.client as win32
import docx.oxml as oxml
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from os import listdir
from os.path import isfile, join, abspath


class Docx:
    def __init__(self, doc_path: str):
        self.logs_folder = "logs/"
        self.output_path = "output/output.docx"
        self.doc = docx.Document(doc_path)
        self.tag = None

    def _fill_logs(self) -> str:
        self.tag = "<LOG>"
        n_table, n_row, n_cell = self.__find_tag_line()

        table = self.doc.tables[n_table]
        cell = table.cell(n_row, n_cell)

        filenames_dict = self.__get_filenames()
        filenames = filenames_dict["filenames_text"].split(',')

        self.__write_cell(cell, filenames_dict["filenames_text"])

        self.__save_docx()

        self.__attach_files(filenames, n_table+1)

        return self.output_path

    def __save_docx(self):
        try:
            self.doc.save(self.output_path)
        except Exception as err:
            raise Exception(f"Error when try to save output [{err}]")

    def __save_win32_changes(self):
        try:
            self.win32_word.ActiveDocument.SaveAs(abspath(self.output_path))
            self.win32_doc.Close()
        except Exception as err:
            raise Exception(f"Error when try to save win32 changes [{err}]")

    def __find_tag_line(self) -> list:
        error_message = f"TAG [{self.tag}] not found"
        tag_coordinates = []

        for table_count, table in enumerate(self.doc.tables):
            for row_count, row in enumerate(table.rows):
                for cell_count, cell in enumerate(row.cells):
                    if self.tag in cell.text:
                        tag_coordinates = [table_count, row_count, cell_count]
                        break

        if not tag_coordinates:
            raise Exception(error_message)

        return tag_coordinates

    @staticmethod
    def __write_cell(cell, text):
        cell.text = text

    def __attach_files(self, filenames, position):
        self.win32_word = win32.gencache.EnsureDispatch('Word.Application')
        self.win32_doc = self.win32_word.Documents.Open(abspath(self.output_path))
        self.win32_doc.Visible = False

        for filename in filenames:
            log_full_path = abspath(f"{self.logs_folder}{filename}")
            self.win32_doc.InlineShapes.AddOLEObject(FileName=log_full_path, Range=self.win32_doc.Tables(position).Cell(0, 0).Range)

        self.__save_win32_changes()

    def __get_filenames(self) -> dict:
        folder = self.logs_folder
        res = {
                    "filenames_text": str,
                    "number_of_filenames": int
                    }

        filenames = [f for f in listdir(folder) if isfile(join(folder, f))]

        res["filenames_text"] = ','.join(filenames)
        res["number_of_filenames"] = len(filenames)

        return res

    @staticmethod
    def __delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    @staticmethod
    def __add_hyperlink(paragraph, text: str, url: str):
        part = paragraph.part
        r_id = part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(oxml.shared.qn('r:id'), r_id, )

        new_run = oxml.shared.OxmlElement('w:r')
        rPr = oxml.shared.OxmlElement('w:rPr')

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        r = paragraph.add_run()
        r._r.append(hyperlink)

        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True
